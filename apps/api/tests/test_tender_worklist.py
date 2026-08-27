"""Тендерный отбор через HTTP.

Ядро здесь не трогается по-настоящему: у него своя база, свои доступы и свои
тесты. Проверяется перевод на язык HTTP — и в первую очередь то, кому что
достаётся. Таблица собирается из колонок книги автоматически, и колонка,
добавленная в ядре, попала бы в ответ сама, вместе с себестоимостью.

Отдельно проверяется, что колонки и значения приходят из одного места. Книга
описывает лист «Отбор» заголовками и функцией значений, и разъехаться они
могут незаметно: значение встанет не под своим заголовком, а заметит это тот,
кто в отчёте увидит маржу в графе количества.
"""

from __future__ import annotations

import unicodedata
from dataclasses import dataclass, field
from decimal import Decimal
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest
from fastapi.testclient import TestClient
from platform_api.db.models import Role
from sqlalchemy.orm import Session as DbSession
from tests.conftest import sign_in


@dataclass
class _Row:
    """Строка книги в том виде, в каком её собирает ядро."""

    title: str = "Насос центробежный ЦНС 60-330"
    customer: str = "АО «Акбастау»"
    ens_code: str = "281312.500.000001"
    average_quote: Decimal | None = Decimal("4761300")
    quantity: Decimal | None = Decimal("56")
    total: Decimal | None = Decimal("266632800")
    cost: Decimal | None = Decimal("123252640")
    subject: str = "Поставка насосов"
    category: str = "Насосное оборудование"
    kind: str = "Товар"
    method: str = "Запрос ценовых предложений"
    margin_percent: Decimal | None = Decimal("53.8")
    review: str = ""
    date: str = "12.03.2026"
    quotes: list[Any] = field(default_factory=list)
    cost_lines: list[Any] = field(default_factory=list)
    findings: list[Any] = field(default_factory=list)
    market: list[Any] = field(default_factory=list)
    folder: str = "тендеры/2 аппак/Каныбек_Насосы 56шт"
    folder_path: str = "/Users/analyst/тендеры/2 аппак/Каныбек_Насосы 56шт"
    by_analog: str = ""
    parts: list[Any] = field(default_factory=list)
    missing: list[Any] = field(default_factory=list)
    cost_low: Decimal | None = None
    cost_high: Decimal | None = None


@dataclass(frozen=True)
class _Verdict:
    label: str = "Брать"
    score: int = 78
    reasons: tuple[str, ...] = ("маржа 53,8% выше порога", "два поставщика на рынке")


@dataclass(frozen=True)
class _Ranked:
    row: _Row
    verdict: _Verdict = _Verdict()


@pytest.fixture
def case_folder(tmp_path: Any) -> Any:
    """Папка закупки с настоящим файлом внутри.

    Настоящим, а не выдуманным: платформа отдаёт содержимое с диска, и
    проверять это подменой файловой системы значит проверять подмену.
    """
    folder = tmp_path / "тендеры" / "аппак" / "Насосы 56шт"
    folder.mkdir(parents=True)
    (folder / "ТЗ насосы.pdf").write_bytes("%PDF-1.7\nтехническое задание\n".encode())
    return folder


@pytest.fixture
def offline_worklist(monkeypatch: pytest.MonkeyPatch) -> _Ranked:
    """Готовый отбор вместо обращения к базе ядра.

    Без подмены тест зависел бы от того, что лежит в базе разработчика:
    сегодня двести шестьдесят закупок, завтра ноль — и падают проверки, к
    которым это отношения не имеет.
    """
    import importlib

    from platform_api.modules.tender import worklist as module

    item = _Ranked(row=_Row())
    monkeypatch.setattr(module, "_ranked", lambda: [item])
    monkeypatch.setattr(module, "case_files", lambda _folder: ())
    router = importlib.import_module("platform_api.modules.tender.router")
    monkeypatch.setattr(router, "worklist", module)
    return item


@pytest.fixture
def with_files(monkeypatch: pytest.MonkeyPatch, offline_worklist: _Ranked, case_folder: Any) -> Any:
    """Тот же отбор, но у закупки есть папка с документами."""
    from platform_api.modules.tender import worklist as module

    offline_worklist.row.folder_path = str(case_folder)
    files = (
        module.CaseFile(
            sha256="a" * 64,
            name="ТЗ насосы.pdf",
            kind="ТЗ",
            size_bytes=42,
            path=case_folder / "ТЗ насосы.pdf",
        ),
        module.CaseFile(
            sha256="b" * 64,
            name="Проект МЗ.docx",
            kind="МЗ",
            size_bytes=1024,
            # Архив не подключён: файл в базе есть, на диске его не видно.
            path=case_folder / "нет-такого.docx",
        ),
    )
    monkeypatch.setattr(module, "case_files", lambda folder: files if folder else ())
    return files


# --- рабочий список --------------------------------------------------------


def test_bez_sessii_ne_otdaem_nichego(client: TestClient) -> None:
    """За этим адресом себестоимость и маржа."""
    assert client.get("/api/tender/worklist").status_code == 401


def test_kolonki_i_znacheniya_prihodyat_iz_knigi(
    client: TestClient, signed_in: Any, offline_worklist: _Ranked
) -> None:
    """Заголовки — те же, что на листе «Отбор», и значения стоят под своими.

    Проверяется не список заголовков, а совпадение с книгой: список пришлось
    бы править вторым движением после каждой правки ядра, и разъехались бы
    они ровно тогда, когда никто не смотрит.
    """
    from tender_analyze.export.ranked import HEADERS, row_values

    data = client.get("/api/tender/worklist").json()

    assert [item["title"] for item in data["columns"]] == list(HEADERS)
    assert data["sheet"] == "Отбор"

    expected = row_values(offline_worklist)  # type: ignore[arg-type]
    cells = data["rows"][0]["cells"]
    assert cells[HEADERS.index("название закупки")]["text"] == str(expected[2])
    assert cells[HEADERS.index("сумма")]["number"] == pytest.approx(266632800.0)
    assert cells[HEADERS.index("себестоимость")]["number"] == pytest.approx(123252640.0)


def test_verdikt_krasit_stroku_i_podpisan_slovami(
    client: TestClient, signed_in: Any, offline_worklist: _Ranked
) -> None:
    """Цвет сам по себе смысла не несёт: при дальтонизме «брать» и «мимо»
    неразличимы, поэтому слово приходит вместе с ним."""
    data = client.get("/api/tender/worklist").json()

    assert data["rows"][0]["tone"] == "good"
    assert ("good", "Брать") in [(item["tone"], item["title"]) for item in data["legend"]]
    # Счётчики ключуются тоном, а не словом: браузер думает тонами, и слово
    # «Брать» в разделе, где вердикты называются иначе, дало бы молчаливый ноль.
    assert data["verdicts"] == {"good": 1}


def test_knopok_obnovit_i_pereschitat_zdes_net(
    client: TestClient, signed_in: Any, offline_worklist: _Ranked
) -> None:
    """Папки разбирают на машине тендерщика. Кнопка, которую нечем обслужить,
    хуже отсутствующей: человек нажимает и получает ошибку."""
    data = client.get("/api/tender/worklist").json()

    assert data["actions"] == ["export"]


def test_zakupshchik_ne_poluchaet_sebestoimost_dazhe_zaprosom(
    db: DbSession, app_client: TestClient, offline_worklist: _Ranked
) -> None:
    """Спрятать колонку в браузере и отдать её в JSON — самый частый способ
    отдать себестоимость наружу. Проверяется именно ответ."""
    import json

    sign_in(db, app_client, Role.BUYER)

    data = app_client.get("/api/tender/worklist").json()

    titles = [item["title"] for item in data["columns"]]
    assert "себестоимость" not in titles
    assert "заработок" not in titles
    assert "моржа %" not in titles
    assert data["hidden_columns"] == 3
    assert data["margin_total"] is None
    # Ни в одной ячейке — ни числом, ни в тексте.
    assert "123252640" not in json.dumps(data, ensure_ascii=False)
    # И книгу с теми же числами закупщику не отдают — ни кнопкой, ни запросом.
    assert data["actions"] == []
    assert app_client.get("/api/tender/worklist/export").status_code == 403


def test_neznakomaya_kolonka_schitaetsya_denezhnoy() -> None:
    """Таблица собирается автоматически, и колонка, добавленная в ядре,
    попала бы в ответ сама. Умолчание должно быть закрытым."""
    from platform_api.modules.table import Visibility, visible_columns
    from platform_api.modules.tender.columns import POLICY

    @dataclass(frozen=True)
    class _Unknown:
        title: str = "закупочная цена поставщика"
        width: int = 20
        number_format: str | None = None
        getter: Any = None
        hyperlink: Any = None

    for role in (Role.BUYER, Role.VIEWER):
        assert visible_columns([_Unknown()], policy=POLICY, role=role) == []
    assert POLICY["себестоимость"] is Visibility.MONEY


# --- разбор одной закупки --------------------------------------------------


def test_razbor_otkryvaetsya_po_ustoychivomu_imeni(
    client: TestClient, signed_in: Any, offline_worklist: _Ranked
) -> None:
    """Имя строки — не её номер: он съезжает от пересортировки, а ссылку на
    разбор пересылают коллеге."""
    from platform_api.modules.tender.worklist import row_id

    data = client.get("/api/tender/worklist").json()
    identifier = data["rows"][0]["id"]
    assert identifier == row_id(offline_worklist)  # type: ignore[arg-type]

    detail = client.get(f"/api/tender/item/{identifier}").json()
    assert detail["verdict"] == "Брать"
    assert detail["tone"] == "good"
    # Ссылки на площадку у тендерной закупки нет: она пришла папкой по почте.
    assert detail["url"] is None
    assert [section["title"] for section in detail["sections"]][:3] == [
        "Закупка",
        "Решение",
        "Деньги",
    ]


def test_neizvestnaya_stroka_eto_404_a_ne_pustoy_razbor(
    client: TestClient, signed_in: Any, offline_worklist: _Ranked
) -> None:
    assert client.get("/api/tender/item/0000000000000000").status_code == 404


def test_zakupshchik_ne_vidit_razdelov_s_dengami(
    db: DbSession, app_client: TestClient, offline_worklist: _Ranked
) -> None:
    """Та же граница, что у колонок, и по той же причине."""
    sign_in(db, app_client, Role.BUYER)

    detail = app_client.get(f"/api/tender/item/{_identifier(offline_worklist)}").json()

    titles = [section["title"] for section in detail["sections"]]
    assert "Деньги" not in titles
    assert "Как считали себестоимость" not in titles
    # Молча урезанный разбор выглядит недоделанным — про скрытое сказано.
    assert detail["hidden_sections"] > 0


def _identifier(item: _Ranked) -> str:
    from platform_api.modules.tender.worklist import row_id

    return row_id(item)  # type: ignore[arg-type]


# --- документы закупки -----------------------------------------------------


def test_gde_kupit_bez_lishnih_kolonok(
    client: TestClient, signed_in: Any, offline_worklist: _Ranked
) -> None:
    """Доставка, срок и минимальная партия убраны: они съедали треть ширины
    ради «1 шт» и «20 дн.», одинаковых почти во всех строках."""
    offline_worklist.row.market = [_Finding()]

    data = client.get(f"/api/tender/item/{_identifier(offline_worklist)}").json()

    section = next(s for s in data["sections"] if s["title"] == "Где купить")
    assert section["table"]["columns"] == [
        "Позиция",
        "Страна",
        "Площадка",
        "Поставщик",
        "Цена, ₸",
        "Тот ли товар",
    ]


def test_dokumenty_zakupki_vidny_spiskom(
    client: TestClient, signed_in: Any, with_files: Any, offline_worklist: _Ranked
) -> None:
    """Последний вопрос перед подачей — «покажи само ТЗ»."""
    item_id = _identifier(offline_worklist)

    data = client.get(f"/api/tender/item/{item_id}").json()

    section = next(s for s in data["sections"] if s["title"] == "Документы закупки")
    # Свёрнут: файлов в папке бывает три десятка.
    assert section["collapsed"] is True
    assert [(f["label"], f["text"]) for f in section["fields"]] == [
        ("ТЗ", "ТЗ насосы.pdf"),
        ("МЗ", "Проект МЗ.docx"),
    ]
    # Что есть на диске — ссылкой, чего нет — словами. Молчание про
    # неподключённый архив читается как «платформа сломалась».
    assert section["fields"][0]["link"] == f"/api/tender/item/{item_id}/file/{'a' * 64}"
    assert section["fields"][1]["link"] is None
    assert section["fields"][1]["note"] == "нет на диске"
    # Сообщение называет путь, по которому искали: «нет на диске» отвечает,
    # что случилось, но не что чинить.
    assert "Искали в" in section["note"]
    assert str(with_files[0].path.parent) in section["note"]


def test_dokument_otdaetsya_s_diska(
    client: TestClient, signed_in: Any, with_files: Any, offline_worklist: _Ranked
) -> None:
    response = client.get(f"/api/tender/item/{_identifier(offline_worklist)}/file/{'a' * 64}")

    assert response.status_code == 200
    assert response.content.startswith(b"%PDF")
    # Открыть, а не скачать: за ТЗ идут посмотреть.
    assert "inline" in response.headers["content-disposition"]


def test_dokument_iz_chuzhoy_zakupki_ne_otdaetsya(
    client: TestClient, signed_in: Any, with_files: Any, offline_worklist: _Ranked
) -> None:
    """Взять файл по одному хэшу значило бы прочитать чужую папку, подставив
    хэш оттуда, — а в чужой папке лежат КП с ценами, которых нам не показывали."""
    assert client.get(f"/api/tender/item/0000000000000000/file/{'a' * 64}").status_code == 404


def test_nesuschestvuyuschiy_fayl_govorit_pochemu(
    client: TestClient, signed_in: Any, with_files: Any, offline_worklist: _Ranked
) -> None:
    """Файл в базе есть, а на диске его нет: это не поломка платформы, а
    неподключённый архив, и сказать это словами полезнее пустого ответа."""
    response = client.get(f"/api/tender/item/{_identifier(offline_worklist)}/file/{'b' * 64}")

    assert response.status_code == 404
    assert "TENDER_ARCHIVE" in response.json()["detail"]


def test_dokumenty_zakryty_bez_sessii(client: TestClient, offline_worklist: _Ranked) -> None:
    assert client.get(f"/api/tender/item/x/file/{'a' * 64}").status_code == 401


@dataclass(frozen=True)
class _Finding:
    """Находка на рынке в том виде, в каком её отдаёт ядро."""

    position: str = "Насос ЦНС 60-330"
    country: str = "Казахстан"
    marketplace: str = "satu.kz"
    supplier: str = "ТОО «KARLSKRONA»"
    price: Decimal = Decimal("3650000")
    landed: Decimal = Decimal("3857500")
    delivery_days: int = 15
    min_order: str = "1 шт"
    matches_spec: bool = True
    note: str = ""
    url: str = "https://satu.kz/p/nasos-tsns"
    """Ссылка на карточку. У настоящей находки она есть, и подделка без неё
    прятала бы поломку: разбор обращается к ней при каждом открытии."""


# --- выбор поставщика ------------------------------------------------------


@dataclass(frozen=True)
class _Country:
    value: str = "Казахстан"


@dataclass(frozen=True)
class _Market:
    """Находка так, как её хранит ядро.

    Отдельно от `_Finding`: та описывает строку книги, а эта — доменный
    объект, и поля у них разные. Одна подделка на две роли скрыла бы, что
    разбор читает то одну, то другую.
    """

    marketplace: str = "dn.ru"
    supplier: str = "ДН.РУ"
    price_kzt: Decimal = Decimal("2000")
    matches_spec: bool = True
    country: _Country = _Country()
    url: str = "https://dn.ru/nasos"


@dataclass(frozen=True)
class _Chosen:
    """Находка вместе с посчитанной по ней себестоимостью — как у ядра."""

    position: str
    landed_cost: Decimal
    quantity: Decimal | None
    finding: _Market


@pytest.fixture
def with_sourcing(monkeypatch: pytest.MonkeyPatch, offline_worklist: _Ranked) -> Any:
    """Две находки по одной позиции: дешёвая непроходная и дорогая подходящая.

    Ровно тот случай, ради которого выбор и сделан: считать по дешёвой нельзя,
    её нельзя поставить, но и разница в деньгах должна быть видна.
    """
    from platform_api.modules.tender import worklist as module

    cheap = _Chosen(
        position="Насос",
        landed_cost=Decimal("1000"),
        quantity=None,
        finding=_Market(
            marketplace="1688.com",
            supplier="CNP",
            price_kzt=Decimal("1000"),
            matches_spec=False,
        ),
    )
    proper = _Chosen(
        position="Насос", landed_cost=Decimal("2000"), quantity=None, finding=_Market()
    )

    @dataclass(frozen=True)
    class _Sourcing:
        opportunities: tuple[_Chosen, ...]

    monkeypatch.setattr(module, "case_sourcing", lambda _folder: (_Sourcing((cheap, proper)), 1))
    # Сумма закупки поменьше: при 266 миллионах обе себестоимости дают маржу
    # 100,0 % после округления, и проверка «маржа пересчиталась» не поймала бы
    # ничего.
    offline_worklist.row.total = Decimal("200000")

    # Строки книги идут в том же порядке и с теми же площадками: по ним разбор
    # и сопоставляет находки со своими ключами.
    offline_worklist.row.market = [
        _Finding(
            position="Насос",
            marketplace="1688.com",
            supplier="CNP",
            matches_spec=False,
            price=Decimal("1000"),
        ),
        _Finding(position="Насос", marketplace="dn.ru", supplier="ДН.РУ", price=Decimal("2000")),
    ]
    return module, cheap, proper


def test_po_umolchaniyu_beretsya_podhodyashchaya_a_ne_deshevaya(
    client: TestClient, signed_in: Any, with_sourcing: Any, offline_worklist: _Ranked
) -> None:
    """Товар, который не проходит по требованиям, поставить нельзя — какой бы
    ни была его цена. Дешёвая находка с пометкой «нет» в расчёт не идёт."""
    module, cheap, proper = with_sourcing

    data = client.get(f"/api/tender/item/{_identifier(offline_worklist)}").json()

    table = next(s for s in data["sections"] if s["title"] == "Где купить")["table"]
    assert table["chosen"] == [module.finding_key(proper)]
    assert module.finding_key(cheap) not in table["chosen"]


def test_vybor_postavshchika_pereschityvaet_dengi(
    client: TestClient, signed_in: Any, with_sourcing: Any, offline_worklist: _Ranked
) -> None:
    """«Подходит» — суждение модели, и тендерщик вправе с ним не согласиться.
    Выбрал другую находку — деньги должны стать её."""
    module, cheap, _proper = with_sourcing
    item_id = _identifier(offline_worklist)

    before = _money(client.get(f"/api/tender/item/{item_id}").json())
    after = _money(
        client.get(f"/api/tender/item/{item_id}?pick={module.finding_key(cheap)}").json()
    )

    # 56 штук по 2000 против 56 по 1000 — себестоимость ровно вдвое меньше.
    assert before["Себестоимость"] == pytest.approx(112000.0)
    assert after["Себестоимость"] == pytest.approx(56000.0)
    # И заработок с маржой пересчитаны под неё, а не остались от умолчания.
    assert (before["Заработок"], before["Маржа"]) == (pytest.approx(88000.0), "44.0 %")
    assert (after["Заработок"], after["Маржа"]) == (pytest.approx(144000.0), "72.0 %")


def test_vybor_otmechen_v_spiske_nahodok(
    client: TestClient, signed_in: Any, with_sourcing: Any, offline_worklist: _Ranked
) -> None:
    """Без отметки непонятно, откуда взялась цифра, и человек считает её
    заново на глаз."""
    module, cheap, _proper = with_sourcing

    data = client.get(
        f"/api/tender/item/{_identifier(offline_worklist)}?pick={module.finding_key(cheap)}"
    ).json()

    table = next(s for s in data["sections"] if s["title"] == "Где купить")["table"]
    assert table["chosen"] == [module.finding_key(cheap)]


def test_nesushchestvuyushchiy_vybor_ne_lomaet_razbor(
    client: TestClient, signed_in: Any, with_sourcing: Any, offline_worklist: _Ranked
) -> None:
    """Ссылку с выбором пересылают, а отбор к тому времени пересобран: находки
    той уже нет. Разбор должен показать умолчание, а не отказать."""
    response = client.get(f"/api/tender/item/{_identifier(offline_worklist)}?pick=неттакой")

    assert response.status_code == 200
    assert _money(response.json())["Себестоимость"] == pytest.approx(112000.0)


def test_nahodki_otdayutsya_so_ssylkami(
    client: TestClient, signed_in: Any, with_sourcing: Any, offline_worklist: _Ranked
) -> None:
    """Находка без ссылки — обещание, а не поставщик: менеджер идёт искать её
    заново поиском и половину не находит."""
    data = client.get(f"/api/tender/item/{_identifier(offline_worklist)}").json()

    table = next(s for s in data["sections"] if s["title"] == "Где купить")["table"]
    assert all(link for link in table["links"])
    # Ссылка на площадке, а не на названии: название в строках одно и то же.
    assert table["columns"][table["link_column"]] == "Площадка"


def _money(detail: dict[str, Any]) -> dict[str, Any]:
    fields = next(s for s in detail["sections"] if s["title"] == "Деньги")["fields"]
    return {
        item["label"]: (item["number"] if item["number"] is not None else item["text"])
        for item in fields
    }


def test_fail_zakupki_nahoditsya_v_lyuboy_forme_zapisi(tmp_path: Path) -> None:
    """Имя на диске и имя в базе могут расходиться формой записи Unicode.

    Разбор шёл на macOS, где поиск файла к этой разнице нечувствителен, а
    сервер на ext4 сравнивает байты. Хуже того, одна и та же папка попадает
    в базу в обоих написаниях — разными прогонами, — поэтому переименовать
    диск под базу нельзя: какую форму ни выбери, вторая половина путей
    перестанет сходиться. Находить файл должна платформа.
    """
    from platform_api.modules.tender.worklist import _LISTINGS, _on_disk

    _LISTINGS.clear()
    folder = tmp_path / unicodedata.normalize("NFD", "Жанайбек_Страховка")
    folder.mkdir()
    (folder / unicodedata.normalize("NFD", "КП ОГПО ёлки.pdf")).write_text("тз", encoding="utf-8")

    for form in ("NFC", "NFD"):
        asked = Path(unicodedata.normalize(form, str(folder / "КП ОГПО ёлки.pdf")))
        found = _on_disk(asked)
        assert found.is_file(), f"не нашёлся по пути в форме {form}"
        assert found.read_text(encoding="utf-8") == "тз"


def test_fail_ne_iz_svoey_papki_ne_otdaetsya() -> None:
    """Приведение имён не должно открывать дорогу за пределы папки закупки."""
    from platform_api.modules.tender.worklist import _inside

    assert not _inside(Path("/srv/tenders/чужая/тз.pdf"), "/srv/tenders/своя")
    assert not _inside(Path("/srv/tenders/своя/../чужая/тз.pdf"), "/srv/tenders/своя")
    assert _inside(Path("/srv/tenders/своя/тз.pdf"), "/srv/tenders/своя")


def test_fail_stroki_odnoy_papki_poluchayut_svoi_dokumenty(tmp_path: Path) -> None:
    """В одной папке лежат бумаги нескольких потребностей.

    Пять служебных записок на пять разных нужд — обычное дело, а строка в
    отборе своя у каждой позиции. Пока файлы висели на папке, все пять строк
    показывали все пять записок: человек открывал документ, читал про чужие
    насосы и переставал верить разбору.
    """
    from platform_api.modules.tender.worklist import _attach_files

    def document(name: str, kind: str) -> Any:
        path = tmp_path / name
        path.write_text("x", encoding="utf-8")
        source = SimpleNamespace(
            sha256=name, name=name, size_bytes=1, path=path, relative_path=Path(name)
        )
        return SimpleNamespace(source=source, insight=SimpleNamespace(kind=kind))

    case = SimpleNamespace(
        documents=[
            document("МЗ насосы.docx", "МЗ"),
            document("МЗ кабель.docx", "МЗ"),
            document("КП насосы.pdf", "КП"),
            document("Справка.pdf", ""),
        ]
    )

    def row(title: str, sources: list[str], quotes: list[str]) -> Any:
        return SimpleNamespace(
            title=title,
            ens_code="",
            folder_path=str(tmp_path),
            sources=sources,
            quotes=[("Поставщик", name, None, "", "") for name in quotes],
        )

    pumps = row("Насосы", ["МЗ насосы.docx"], ["КП насосы.pdf"])
    cable = row("Кабель", ["МЗ кабель.docx"], [])

    files: dict[str, Any] = {}
    _attach_files(files, [pumps, cable], case)

    from platform_api.modules.tender.worklist import row_id_of_folder

    seen = {key: [(item.name, item.shared) for item in value] for key, value in files.items()}
    assert seen[row_id_of_folder(pumps)] == [
        ("МЗ насосы.docx", False),
        ("КП насосы.pdf", False),
        # Справка про закупку целиком — достаётся обеим строкам, но помечена:
        # прятать её было бы враньём наоборот.
        ("Справка.pdf", True),
    ]
    assert seen[row_id_of_folder(cable)] == [
        ("МЗ кабель.docx", False),
        ("Справка.pdf", True),
    ]


def test_fail_odna_stroka_na_papku_vidit_vsyu_papku(tmp_path: Path) -> None:
    """Делить не с кем — значит всё в папке принадлежит единственной строке."""
    from platform_api.modules.tender.worklist import _attach_files, row_id_of_folder

    path = tmp_path / "ТЗ.pdf"
    path.write_text("x", encoding="utf-8")
    case = SimpleNamespace(
        documents=[
            SimpleNamespace(
                source=SimpleNamespace(
                    sha256="a", name="ТЗ.pdf", size_bytes=1, path=path, relative_path=Path("ТЗ.pdf")
                ),
                insight=SimpleNamespace(kind="ТЗ"),
            )
        ]
    )
    only = SimpleNamespace(
        title="Насосы", ens_code="", folder_path=str(tmp_path), sources=[], quotes=[]
    )

    files: dict[str, Any] = {}
    _attach_files(files, [only], case)

    assert [item.name for item in files[row_id_of_folder(only)]] == ["ТЗ.pdf"]


def test_fail_staroe_yadro_ne_ostavlyaet_razdel_pustym(tmp_path: Path) -> None:
    """Платформа и ядро — разные репозитории, выкатываются порознь.

    Строка из старого ядра не знает, из какого документа взята позиция.
    Разложить по одним предложениям нельзя — строка осталась бы без своего
    ТЗ, — но и падать нельзя: раздел целиком отвечал «Данные недоступны», и
    разбор было не открыть вовсе.
    """
    from platform_api.modules.tender.worklist import _attach_files, row_id_of_folder

    path = tmp_path / "ТЗ.pdf"
    path.write_text("x", encoding="utf-8")
    case = SimpleNamespace(
        documents=[
            SimpleNamespace(
                source=SimpleNamespace(
                    sha256="a", name="ТЗ.pdf", size_bytes=1, path=path, relative_path=Path("ТЗ.pdf")
                ),
                insight=SimpleNamespace(kind="ТЗ"),
            )
        ]
    )
    # Ровно старая строка: `quotes` есть, `sources` ещё нет.
    old = [
        SimpleNamespace(title=name, ens_code="", folder_path=str(tmp_path), quotes=[])
        for name in ("Насосы", "Кабель")
    ]

    files: dict[str, Any] = {}
    _attach_files(files, old, case)

    for row in old:
        seen = files[row_id_of_folder(row)]
        assert [item.name for item in seen] == ["ТЗ.pdf"]
        # Помечены общими, и это правда: чей документ — сказать нечем.
        assert all(item.shared for item in seen)


def test_fail_kod_iz_perechnya_krasit_svoyu_yacheyku() -> None:
    """Перечень Минпрома отмечает код, а не строку.

    Заливка строки занята вердиктом — «брать», «мимо», — и второго смысла в
    неё не вложить: закупка бывает и выгодной, и с отечественным
    производителем разом. Красится сам код: он стоит первой колонкой, его
    видно, не читая остального, а вердикт остаётся при строке.
    """
    import platform_api.modules.tender.worklist as module

    own = SimpleNamespace(row=SimpleNamespace(ens_code="281314.900.000076"))
    foreign = SimpleNamespace(row=SimpleNamespace(ens_code="999999.999.999999"))

    assert module.mark_cell("ЕНС ТРУ", own) == "critical"
    # Другие колонки той же строки не трогаются.
    assert module.mark_cell("сумма", own) == ""
    assert module.mark_cell("ЕНС ТРУ", foreign) == ""
    # И вердикт строки остаётся вердиктом.
    assert (
        module.tone_of(SimpleNamespace(row=own.row, verdict=SimpleNamespace(label="Брать")))
        == "good"
    )


def test_fail_lot_sobiraet_sosednie_pozicii_i_ih_itog() -> None:
    """Заработок по одной позиции ничего не значит без остальных.

    В заключении три позиции, по одной из них маржа отличная — её берут в
    работу, а поставить придётся все три, и на остальных убыток. В сумме
    сделка убыточна, и увидеть это надо до подачи.
    """
    from platform_api.modules.tender.lots import collect

    def line(sheet_name: str, amount: str, cost_value: str, margin: str) -> Any:
        return SimpleNamespace(
            row=SimpleNamespace(
                folder_path="/архив/Закупка",
                title=sheet_name,
                ens_code="",
                quantity=Decimal(1),
                total=Decimal(amount),
                cost=Decimal(cost_value),
                margin_percent=Decimal(margin),
            ),
            verdict=SimpleNamespace(label="Брать"),
        )

    profitable = line("Насос", "1000000", "600000", "40.0")
    rows = [
        profitable,
        line("Кабель", "500000", "700000", "-40.0"),
        line("Щит", "300000", "500000", "-66.7"),
    ]

    lot = collect(rows, profitable, money=True)

    assert lot is not None and lot.size == 3
    assert lot.total == Decimal(1_800_000)
    assert lot.cost == Decimal(1_800_000)
    # Позиция даёт 40%, а лот целиком — ноль: вот ради чего он и собирается.
    assert lot.profit == Decimal(0)
    assert lot.margin_percent == Decimal("0.0")
    assert [p.current for p in lot.positions] == [True, False, False]


def test_fail_lot_ne_otdaet_dengi_zakupshchiku() -> None:
    """Итог по лоту — та же себестоимость, только сложенная.

    Отдать её потому, что она лежит в другом поле, значило бы обойти
    собственные права: закупщик не видит себестоимость по строке.
    """
    from platform_api.modules.tender.lots import collect

    def line(sheet_name: str) -> Any:
        return SimpleNamespace(
            row=SimpleNamespace(
                folder_path="/архив/Закупка",
                title=sheet_name,
                ens_code="",
                quantity=Decimal(1),
                total=Decimal(1000),
                cost=Decimal(600),
                margin_percent=Decimal("40.0"),
            ),
            verdict=SimpleNamespace(label="Брать"),
        )

    rows = [line("Насос"), line("Кабель")]
    lot = collect(rows, rows[0], money=False)

    assert lot is not None
    assert lot.total == Decimal(2000), "сумма закупки — не тайна"
    assert lot.cost is None and lot.profit is None and lot.margin_percent is None
    assert all(p.cost is None and p.margin_percent is None for p in lot.positions)


def test_lot_iz_odnoy_pozicii_ne_lot() -> None:
    """Объединять нечего — и кнопки быть не должно."""
    from platform_api.modules.tender.lots import collect

    single = SimpleNamespace(
        row=SimpleNamespace(
            folder_path="/архив/Закупка",
            title="Насос",
            ens_code="",
            quantity=Decimal(1),
            total=Decimal(1000),
            cost=Decimal(600),
            margin_percent=Decimal("40.0"),
        ),
        verdict=SimpleNamespace(label="Брать"),
    )

    assert collect([single], single, money=True) is None


def test_fail_sostav_lota_silnee_papki() -> None:
    """Разбор связывает позиции папкой, и иногда ошибается.

    Заказчик раскладывает один лот по двум папкам, и признака этого в
    документах нет. Утверждённый человеком состав должен побеждать догадку:
    иначе исправить ошибку нечем.
    """
    from platform_api.modules.tender.lots import collect

    def line(dir_path: str, sheet_name: str, amount: str) -> Any:
        return SimpleNamespace(
            row=SimpleNamespace(
                folder_path=dir_path,
                title=sheet_name,
                ens_code="",
                quantity=Decimal(1),
                total=Decimal(amount),
                cost=Decimal(amount) / 2,
                margin_percent=Decimal("50.0"),
            ),
            verdict=SimpleNamespace(label="Брать"),
        )

    rows = [
        line("/архив/А", "Насос", "1000"),
        line("/архив/А", "Кабель", "2000"),
        line("/архив/Б", "Щит", "4000"),
    ]

    # Без состава — подсказка по папке: только соседи «Насоса».
    hint = collect(rows, rows[0], money=True)
    assert hint is not None and hint.merged is False
    assert [p.title for p in hint.positions] == ["Насос", "Кабель"]

    # С составом — ровно перечисленные, в том числе из другой папки.
    own = collect(
        rows,
        rows[0],
        members=frozenset({("/архив/А", "Насос"), ("/архив/Б", "Щит")}),
        key="лот-1",
        money=True,
    )
    assert own is not None and own.merged is True and own.key == "лот-1"
    assert [p.title for p in own.positions] == ["Насос", "Щит"]
    assert own.total == Decimal(5000)


def test_fail_stroki_lota_stoyat_ryadom() -> None:
    """Позиция, добавленная вручную, должна встать к своим.

    Порядок задаёт ядро — по выгоде, — и позиции лота он разбрасывает по всему
    списку: добавленная из чужой папки оказывается через двести строк от
    своих. Связь тогда видна только по значку, а рядом её нет.
    """
    from platform_api.modules.tender.router import _grouped

    def line(dir_path: str, sheet_name: str) -> Any:
        return SimpleNamespace(row=SimpleNamespace(folder_path=dir_path, title=sheet_name))

    rows = [
        line("/а", "Насос"),
        line("/б", "Кабель"),
        line("/в", "Щит"),
        line("/г", "Лоток"),
    ]
    # В лоте первая и последняя — между ними двести строк в жизни.
    marked = {("/а", "Насос"): "лот-1", ("/г", "Лоток"): "лот-1"}

    ordered = [item.row.title for item in _grouped(rows, marked, {})]

    # Лот встал на место лучшей своей позиции, остальные не сдвинулись.
    assert ordered == ["Насос", "Лоток", "Кабель", "Щит"]
    # Разъединили — порядок ядра вернулся сам, ничего не запоминается.
    assert [item.row.title for item in _grouped(rows, {}, {})] == [
        "Насос",
        "Кабель",
        "Щит",
        "Лоток",
    ]


def test_fail_kod_stroki_ne_menyaetsya_ot_novyh_zakupok() -> None:
    """Код выдаётся один раз и остаётся при позиции.

    Порядковый номер сдвигается от каждой новой закупки: сотрудник говорит
    «посмотри сорок вторую», а у собеседника это уже другая строка.
    """
    from platform_api.modules import codes

    class Storage:
        """База в памяти: проверяется правило выдачи, а не SQL."""

        def __init__(self) -> None:
            self.issued: dict[tuple[str, str], int] = {}

        def assign(self, module: str, prefix: str, keys: list[str]) -> dict[str, str]:
            next_one = max((n for (m, _k), n in self.issued.items() if m == module), default=0)
            for key in keys:
                if (module, key) not in self.issued:
                    next_one += 1
                    self.issued[(module, key)] = next_one
            return {key: f"{prefix}-{self.issued[(module, key)]:0{codes.WIDTH}d}" for key in keys}

    warehouse = Storage()
    first = warehouse.assign("tender", "TN", ["а", "б", "в"])
    assert first == {"а": "TN-00001", "б": "TN-00002", "в": "TN-00003"}

    # Новая закупка встала первой в списке — коды прежних не тронуты.
    later = warehouse.assign("tender", "TN", ["новая", "а", "б", "в"])
    assert later["а"] == "TN-00001" and later["в"] == "TN-00003"
    assert later["новая"] == "TN-00004"


def test_fail_dokument_word_razbiraetsya_na_abzacy_i_tablicy(tmp_path: Path) -> None:
    """Word и Excel должны открываться в платформе, а не уезжать в загрузки.

    Скачанный файл открывается чужой программой, и обратно к строке человек
    возвращается руками. За смену таких выходов десятки — ТЗ смотрят по
    каждой закупке.
    """
    from docx import Document as WordDocument
    from platform_api.modules import preview

    file = tmp_path / "ТЗ.docx"
    doc = WordDocument()
    doc.add_heading("Техническое задание", level=1)
    doc.add_paragraph("Пикобур 3-х лопастной с резцами PDC Ø190,5")
    grid = doc.add_table(rows=2, cols=2)
    grid.cell(0, 0).text = "Наименование"
    grid.cell(0, 1).text = "Кол-во"
    grid.cell(1, 0).text = "Пикобур"
    grid.cell(1, 1).text = "105"
    doc.save(file)

    parsed = preview.build(file)

    assert parsed.kind == "document"
    kinds = [block.kind for block in parsed.blocks]
    # По порядку, как в файле: таблица идёт за своим заголовком, и разложенные
    # порознь они теряют смысл.
    assert kinds == ["heading", "text", "table"]
    assert parsed.blocks[2].rows[0] == ("Наименование", "Кол-во")


def test_fail_neizvestnyy_format_govorit_pochemu(tmp_path: Path) -> None:
    """Старые «.doc» платформа не читает — и должна сказать это словами.

    Пустое окно читается как поломка платформы, а не как «этот формат мы не
    умеем».
    """
    from platform_api.modules import preview

    file = tmp_path / "Старое ТЗ.doc"
    file.write_bytes(b"\xd0\xcf\x11\xe0")

    parsed = preview.build(file)

    assert parsed.kind == "none"
    assert "doc" in parsed.note and "Word" in parsed.note


def test_fail_bityy_fayl_ne_ronyaet_prosmotr(tmp_path: Path) -> None:
    """Битый файл — повод сказать и предложить скачать, а не отдать пятисотую."""
    from platform_api.modules import preview

    file = tmp_path / "Обрезано.docx"
    file.write_bytes(b"PK\x03\x04 truncated")

    parsed = preview.build(file)

    assert parsed.kind == "none"
    assert "повреждён" in parsed.note


# --- лот в работе между отделами -------------------------------------------


@pytest.fixture
def organization(db: Any) -> Any:
    """Организация, которой принадлежит работа."""
    import uuid as _uuid

    from platform_api.db.models import Organization

    org = Organization(name="Fintend", slug=f"fintend-{_uuid.uuid4().hex[:6]}")
    db.add(org)
    db.flush()
    return org


def _work(db: Any, org: Any, how_many: int = 2) -> Any:
    from platform_api.modules.tender import works

    return works.take(
        db,
        org.id,
        None,
        code="TN-00042",
        title="Пикобуры",
        customer="АО «Волковгеология»",
        positions=[
            works.Draft(
                folder_path="/архив/Пикобуры",
                title=f"Пикобур Ø{190 + n}",
                code=f"TN-0004{n}",
                quantity=Decimal(5),
                unit="шт",
                total=Decimal(1000),
                spec="ТЕХНИЧЕСКОЕ ЗАДАНИЕ\n\nПикобур PDC, 3 лопасти",
                spec_source="ТЗ.pdf",
                options=(
                    {"name": "Пикобур", "supplier": "Первый", "price": Decimal(100)},
                    {"name": "Пикобур", "supplier": "Второй", "price": Decimal(90)},
                ),
            )
            for n in range(how_many)
        ],
    )


def test_fail_zakaz_poiska_ubiraet_otvergnutye_nahodki(db: Any, organization: Any) -> None:
    """«Найденное не подходит, поищите сами» — и найденное уходит.

    Иначе снабжение потратит день на то, что разбор уже посмотрел и отверг, —
    а ради того, чтобы этого не было, процесс и заводили.
    """
    from platform_api.modules.tender import works

    work = _work(db, organization)
    position_value = work.positions[0]
    assert len(position_value.options) == 2

    works.ask(db, work, position_value.id, "Пикобур PDC Ø215, 3 лопасти")
    db.refresh(position_value)

    left = [option.source.value for option in position_value.options]
    assert left == ["asked"]


def test_fail_lot_ne_uhodit_s_nemoy_poziciey(db: Any, organization: Any) -> None:
    """Позиция без единого варианта — это молчание.

    Снабжение получит строку и не поймёт, искать по ней или она попала
    случайно. Спросить об этом оно сможет только письмом, а ради письма
    процесс и заводили.
    """
    import pytest as _pytest
    from platform_api.errors import SpokenError
    from platform_api.modules.tender import works

    work = _work(db, organization)
    for option in list(work.positions[1].options):
        works.drop(db, work, option.id)
    db.refresh(work)

    with _pytest.raises(SpokenError, match="не выбран поставщик"):
        works.hand_over(db, work, "")


def test_fail_snabzhenie_ne_vidit_summ_zakupki(db: Any, organization: Any) -> None:
    """Снабжению суммы не уходят даже в ответе.

    По ним считают маржу, а это не его работа и не его сведения. Спрятать в
    браузере и отдать в JSON — значит отдать.
    """
    from platform_api.db.models import Role
    from platform_api.modules.tender import works
    from platform_api.modules.tender.works_router import _work_out

    work = _work(db, organization)
    works.choose(db, work, work.positions[0].options[0].id)
    works.choose(db, work, work.positions[1].options[0].id)
    works.hand_over(db, work, "проверьте")

    for_supply = _work_out(work, Role.BUYER)
    assert for_supply.total is None and for_supply.cost is None
    assert all(position.total is None for position in for_supply.positions)
    # А «где купить» — уходит: без него работать нечем.
    assert all(position.options for position in for_supply.positions)

    for_analysis = _work_out(work, Role.ANALYST)
    assert for_analysis.total is not None and for_analysis.cost is not None


# ---------------------------------------------------------------------------
# Техническое задание позиции
# ---------------------------------------------------------------------------


def _case_spec() -> Any:
    """Требования закупки, собранные из разобранных ядром документов.

    Через `gather`, а не руками: белый список полей и отсев денег живут именно
    там, и собранный в обход них образец проверял бы не тот код.
    """
    from types import SimpleNamespace

    from platform_api.modules.tender import spec

    def page(text: str, tables: Any = ()) -> Any:
        return SimpleNamespace(
            text=text,
            tables=tuple(SimpleNamespace(rows=rows, is_meaningful=True) for rows in tables),
        )

    def doc(kind: str, name: str, **extra: Any) -> Any:
        return SimpleNamespace(
            source=SimpleNamespace(name=name),
            extraction=SimpleNamespace(pages=extra.get("pages", ())),
            insight=SimpleNamespace(
                kind=kind,
                requirements=extra.get("requirements", []),
                delivery_terms=extra.get("delivery_terms"),
                warranty=extra.get("warranty"),
            ),
        )

    def position_value(name: str, spec_text: str, count: int, ens: str) -> Any:
        return SimpleNamespace(
            name=name, specification=spec_text, quantity=Decimal(count), unit="шт", ens_code=ens
        )

    case = SimpleNamespace(
        subject="Пикобуры PDC",
        requirements=(),
        documents=(
            doc(
                "ТЗ",
                "ТЗ Пикобуры.pdf",
                requirements=[
                    "Наработка на отказ не менее 8000 часов",
                    "Сертификат качества на изделие",
                    "Цена не выше 350 000 тенге за штуку",
                ],
                delivery_terms="DDP, рудник «Западный Мынкудук»",
                warranty="не менее 12 месяцев",
            ),
            # То же требование в МЗ — слово в слово, как это и бывает.
            doc("МЗ", "МЗ.docx", requirements=["Наработка на отказ не менее 8000 часов"]),
            # КП поставщика: его условия требованиями закупки не являются.
            doc("КП", "КП Сервис-А.pdf", requirements=["Оплата 100% предоплата"]),
        ),
        requested=(
            position_value(
                "Пикобур PDC Ø190,5", "3 лопасти, резцы PDC, замок З-88", 105, "289221.500.000015"
            ),
            position_value("Пикобур PDC Ø215", "3 лопасти", 5, "289221.500.000016"),
        ),
    )
    return spec.gather(case)


def test_fail_zadanie_sobirayetsya_po_svoey_pozicii() -> None:
    """Задание собирается по своей позиции, а не по всей папке.

    В папке пять служебных записок на пять разных нужд. Задание, в котором
    перечислены все пять, снабжение читает как «купите всё» — и покупает.
    """
    from platform_api.modules.tender import spec

    body = spec.render(_case_spec(), "Пикобур PDC Ø190,5")

    assert "Пикобур PDC Ø190,5" in body
    assert "Пикобур PDC Ø215" not in body
    assert "Количество: 105 шт" in body
    assert "289221.500.000015" in body
    assert "замок З-88" in body
    # Требования у закупки общие — они и должны быть в задании каждой позиции.
    assert "8000 часов" in body
    # А место поставки — нет: по руднику поставщик найдёт закупку в реестре.
    assert "Мынкудук" not in body and "рудник" not in body


def test_fail_v_zadanie_ne_popadayut_dengi() -> None:
    """Денежное условие в задание не уходит.

    «Цена не выше стольких-то» — это наш потолок, а не свойство товара.
    Ушедшая снабжению цена заказчика торгует против нас: поставщик узнаёт её
    в первом же письме.
    """
    from platform_api.modules.tender import spec

    body = spec.render(_case_spec(), "Пикобур PDC Ø190,5")

    assert "350 000" not in body and "тенге" not in body
    # Условия из чужого КП — тоже не требования закупки: там поставщик пишет
    # то, что удобно ему, и снабжение приняло бы это за требование заказчика.
    assert "предоплата" not in body.casefold()
    # Техническое требование при этом остаётся.
    assert "Сертификат качества" in body


def test_fail_povtoryayushcheesya_trebovanie_odno() -> None:
    """Одно и то же требование печатается один раз.

    Оно дословно повторяется в ТЗ и МЗ, и без сведения задание начинается с
    трёх одинаковых абзацев — их перестают читать вместе со всем остальным.
    """
    from platform_api.modules.tender import spec

    body = spec.render(_case_spec(), "Пикобур PDC Ø190,5")

    assert body.count("Наработка на отказ не менее 8000 часов") == 1


def test_fail_odna_poziciya_v_papke_beryotsya_bez_sravneniya() -> None:
    """Закупка одной позиции: у строки название всей закупки, а не позиции.

    Сравнение имён здесь не сойдётся никогда, и задание вышло бы пустым при
    полностью разобранной закупке.
    """
    from dataclasses import replace

    from platform_api.modules.tender import spec

    whole = _case_spec()
    single = replace(whole, positions=whole.positions[:1])

    body = spec.render(single, "Закупка пикобуров для нужд рудника")

    assert "Пикобур PDC Ø190,5" in body and "Количество: 105 шт" in body


def test_fail_snabzhenie_ne_vidit_ishodnyh_dokumentov(db: Any, organization: Any) -> None:
    """Снабжению уходит задание, а исходные бумаги — нет.

    В ТЗ заказчика стоят цены ценового заключения, реквизиты и печати. Увидев
    цену заказчика, снабженец знает потолок, по которому с ним же и будут
    торговаться поставщики.
    """
    from platform_api.db.models import Role
    from platform_api.modules.tender.works_router import _work_out

    work = _work(db, organization)

    for_supply = _work_out(work, Role.BUYER)
    assert all(not position.documents for position in for_supply.positions)
    assert all("ТЕХНИЧЕСКОЕ ЗАДАНИЕ" in position.spec for position in for_supply.positions)


def test_fail_lot_bez_zadaniya_ne_peredayotsya(db: Any, organization: Any) -> None:
    """Позиция без задания не уезжает в снабжение.

    Исходных документов там не будет, вариант может быть один и без описания —
    снабжение получит строку и не поймёт, что искать. Спросить оно сможет
    только письмом, а ради письма процесс и заводили.
    """
    import pytest as _pytest
    from platform_api.errors import SpokenError
    from platform_api.modules.tender import works

    work = _work(db, organization)
    works.choose(db, work, work.positions[0].options[0].id)
    works.choose(db, work, work.positions[1].options[0].id)
    work.positions[1].spec = "   "

    with _pytest.raises(SpokenError, match="пустое техническое задание"):
        works.hand_over(db, work, "")

    works.set_spec(db, work, work.positions[1].id, "ТЕХНИЧЕСКОЕ ЗАДАНИЕ\n\nПикобур Ø215")
    assert works.hand_over(db, work, "").stage.value == "supply"


def test_fail_zadanie_pravit_tolko_razbor(db: Any, organization: Any) -> None:
    """Снабжение задание читает, но не правит.

    Иначе исчезает единственное место, где записано, что именно просили
    купить, — и отчитывается снабжение по условиям, которые само же и внесло.
    """
    import pytest as _pytest
    from platform_api.errors import SpokenError
    from platform_api.modules.tender import works

    work = _work(db, organization)
    works.choose(db, work, work.positions[0].options[0].id)
    works.choose(db, work, work.positions[1].options[0].id)
    works.hand_over(db, work, "")

    with _pytest.raises(SpokenError, match="отдел разбора"):
        works.set_spec(db, work, work.positions[0].id, "куплю что подешевле")


def test_fail_zadanie_faylom_otkryvayetsya(db: Any, organization: Any) -> None:
    """Файл задания — настоящий .docx, а не текст с расширением.

    Снабжение пересылает его поставщику; файл, который не открылся, означает
    письмо «пришлите нормальный» и потерянный день.
    """
    from io import BytesIO

    from docx import Document
    from platform_api.modules.tender import spec

    work = _work(db, organization)
    position_value = work.positions[0]

    file = Document(
        BytesIO(spec.document(position_value.code, [(position_value.title, position_value.spec)]))
    )
    body = "\n".join(paragraph.text for paragraph in file.paragraphs)
    assert position_value.code in body and "ТЕХНИЧЕСКОЕ ЗАДАНИЕ" in body


def test_fail_zadanie_otdayotsya_faylom_cherez_endpoint(client: Any, db: Any) -> None:
    """Задание скачивается файлом, и имя файла не ломает заголовок.

    Имя русское, а в заголовке ответа кириллице места нет — без кодирования
    браузер получает битую строку и сохраняет файл как «download».
    """
    from platform_api.db.models import Role
    from tests.conftest import sign_in

    org = sign_in(db, client, Role.ANALYST)
    work = _work(db, org)
    db.commit()

    position_value = work.positions[0]
    reply = client.get(f"/api/tender/works/{work.id}/positions/{position_value.id}/spec.docx")

    assert reply.status_code == 200
    assert reply.headers["content-type"].endswith("wordprocessingml.document")
    assert "filename*=UTF-8''" in reply.headers["content-disposition"]
    assert reply.content[:2] == b"PK"

    # Позиции без задания файла нет — пустой .docx хуже отсутствующего:
    # снабжение решит, что требований к товару нет.
    from platform_api.modules.tender import works

    works.set_spec(db, work, position_value.id, "")
    db.commit()
    assert (
        client.get(
            f"/api/tender/works/{work.id}/positions/{position_value.id}/spec.docx"
        ).status_code
        == 404
    )


def test_fail_vernuvshiysya_lot_snova_stol_razbora(db: Any, organization: Any) -> None:
    """Возврат от снабжения — это стол разбора, а не готовый результат.

    Снабжение прислало цены; разбор их читает и вполне может передумать:
    подтверждённый поставщик оказался дороже соседнего или не успевает к сроку.
    Закрытый на правку возвращённый лот означал бы, что передумать можно только
    заведя второй лот по тем же позициям.
    """
    from platform_api.modules.tender import works

    work = _work(db, organization)
    works.choose(db, work, work.positions[0].options[0].id)
    works.choose(db, work, work.positions[1].options[0].id)
    works.hand_over(db, work, "проверьте")
    works.hand_over(db, work, "цены подтвердили")
    assert work.stage.value == "returned"

    # Разбор передумал: берём второго поставщика и правим задание.
    works.set_spec(db, work, work.positions[0].id, "ТЗ\n\nПикобур, уточнено")
    works.ask(db, work, work.positions[1].id, "Пикобур Ø215, поищите ещё")
    assert works.hand_over(db, work, "ещё раз").stage.value == "supply"


def _case_with_documents(*documents: Any) -> Any:
    """Закупка из готовых документов — для проверок самой сборки текста."""
    from types import SimpleNamespace

    return SimpleNamespace(
        subject="Насосы",
        requirements=(),
        documents=documents,
        requested=(
            SimpleNamespace(
                name="Насос ЦНС-60",
                specification=None,
                quantity=Decimal(3),
                unit="шт",
                ens_code="281314.900.000076",
            ),
        ),
    )


def _document(kind: str, name: str, text: str = "", tables: Any = ()) -> Any:
    from types import SimpleNamespace

    return SimpleNamespace(
        source=SimpleNamespace(name=name),
        extraction=SimpleNamespace(
            pages=(
                SimpleNamespace(
                    text=text,
                    tables=tuple(SimpleNamespace(rows=rows, is_meaningful=True) for rows in tables),
                ),
            )
        ),
        insight=SimpleNamespace(kind=kind, requirements=[], delivery_terms=None, warranty=None),
    )


def test_fail_zadanie_beryotsya_iz_teksta_tz() -> None:
    """Задание — текст самого ТЗ, а не пересказ его моделью.

    Снабжение по этому тексту закупает: марка стали, допуск и номер ГОСТа
    должны стоять теми же словами, что у заказчика. Выжимка в три строки
    теряет как раз таблицу с размерами — и приходит не тот насос.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ насосы.pdf",
            text="Материал корпуса: сталь 1.4401 (316 AISI).\nСтепень защиты: IP 68.",
            tables=[(("Параметр", "Значение"), ("Напор", "150 м"))],
        )
    )
    body = spec.render(spec.gather(case), "Насос ЦНС-60")

    assert "1.4401 (316 AISI)" in body and "IP 68" in body
    # Таблицы наравне с текстом: в ТЗ размеры и допуски стоят именно в них.
    assert "Напор · 150 м" in body
    # Имя файла не печатается: «ТЗ ПНА. Западный Мынкудук.pdf» называет рудник.
    assert "ТЗ насосы.pdf" not in body


def test_fail_bez_tz_beryotsya_mz() -> None:
    """Нет технического задания — берём маркетинговое заключение.

    Закупка без ТЗ — обычное дело, и это не повод оставить снабжение без
    описания товара.
    """
    from platform_api.modules.tender import spec

    report_only = _case_with_documents(
        _document("КП", "КП поставщика.pdf", text="Предлагаем насос по выгодной цене"),
        _document("МЗ", "МЗ.docx", text="Насос центробежный, напор не менее 150 м"),
    )
    body = spec.render(spec.gather(report_only), "Насос ЦНС-60")

    assert "напор не менее 150 м" in body
    # Текст поставщика в задание не идёт: там написано то, что удобно ему.
    assert "выгодной цене" not in body


def test_fail_ceny_iz_marketingovogo_zaklyucheniya_ne_uhodyat() -> None:
    """Строка цен из заключения снабжению не уходит.

    Ценовое заключение — документ про цены: в его главной таблице стоит, во
    сколько заказчик оценил закупку и почём предлагали соседи. Одна такая
    строка сообщает поставщику наш потолок раньше, чем начнётся торг.

    Строка позиции при этом остаётся: там количество «5,00», а не сумма
    «1 000,00», и по форме записи одно от другого отличается.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "МЗ",
            "МЗ.docx",
            text="Расчёт маркетинговой цены\n1 Насос ЦНС-60 785 000,00 1 930 000,00",
            tables=[
                (
                    ("№", "Наименование", "Код ЕНС ТРУ", "Ед.", "Кол-во", "Доставка"),
                    ("1", "Насос ЦНС-60", "281314.900.000076", "шт.", "5,00", "DDP"),
                    ("1", "Насос ЦНС-60", "1 000,00", "1 500,00", "1 266,67", ""),
                )
            ],
        )
    )
    body = spec.render(spec.gather(case), "Насос ЦНС-60")

    assert "785 000,00" not in body and "1 000,00" not in body
    # А строка позиции — с количеством, доставкой и кодом — на месте.
    assert "281314.900.000076 · шт. · 5,00 · DDP" in body


def test_fail_pechati_i_rekvizity_ne_uhodyat() -> None:
    """Место подписи, печать и реквизиты в задание не попадают.

    Задание снабжение пересылает поставщику. Бланк заказчика с его подписной
    строкой в этом письме выглядит документом, которого мы не подписывали.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ.pdf",
            text=(
                "Насос центробежный, напор 150 м\n"
                "БИН 123456789012\n"
                "М.П.\n"
                "«___» _______________2026 г.\n"
                "Подпись уполномоченного лица"
            ),
        )
    )
    body = spec.render(spec.gather(case), "Насос ЦНС-60")

    assert "напор 150 м" in body
    for junk in ("БИН", "М.П.", "Подпись", "2026 г."):
        assert junk not in body


def test_fail_tehnicheskoye_chislo_ne_prinimayut_za_summu() -> None:
    """Число с единицей измерения — не сумма.

    «Цена деления 1 г/см³» это характеристика ареометра, «не менее 8000 часов»
    — наработка на отказ. Выброшенные вместе с ценами, они означают закупку не
    того прибора, и молча.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ.pdf",
            text=(
                "Цена деления 1 г/см3, погрешность ±1 г/см3\n"
                "Наработка на отказ не менее 8000 часов\n"
                "Диапазон измерения от 1180 до 1240 кг/м3"
            ),
        )
    )
    body = spec.render(spec.gather(case), "Насос ЦНС-60")

    assert "Цена деления 1 г/см3" in body
    assert "8000 часов" in body and "от 1180 до 1240" in body


def test_fail_dlinnoye_zadaniye_obrezayetsya_s_pometkoy() -> None:
    """Длинное задание обрезается, и об этом сказано.

    Молча обрезанное хуже длинного: снабжение не узнает, что прочло половину,
    и закупит по половине.
    """
    from platform_api.modules.tender import spec

    long_body = "\n".join(f"Пункт {n}: требование к материалу корпуса" for n in range(4000))
    case = _case_with_documents(_document("ТЗ", "ТЗ.pdf", text=long_body))
    body = spec.render(spec.gather(case), "Насос ЦНС-60")

    assert len(body) > spec.MAX_LENGTH
    assert "Документ длиннее" in body


def test_fail_zametka_pishetsya_tem_u_kogo_lot(client: Any, db: Any) -> None:
    """Заметку по позиции пишет тот отдел, у которого лот сейчас.

    Объясняет решение тот, кто его принял. Чужая правка в этой строке означает,
    что снабжение отчитывается словами, которых не писало.
    """
    from platform_api.db.models import Role
    from tests.conftest import sign_in

    org = sign_in(db, client, Role.BUYER)
    work = _work(db, org)
    db.commit()
    position_value = work.positions[0]
    address = f"/api/tender/works/{work.id}/positions/{position_value.id}/note"

    # Лот у разбора — снабжение его вообще не видит.
    assert client.patch(address, json={"note": "взяли дороже"}).status_code == 404

    from platform_api.modules.tender import works

    works.choose(db, work, position_value.id and work.positions[0].options[0].id)
    works.choose(db, work, work.positions[1].options[0].id)
    works.hand_over(db, work, "")
    db.commit()

    reply = client.patch(address, json={"note": "  дороже,\n  но успеет к сроку  "})
    assert reply.status_code == 200
    # Одна строка: перевод строки и лишние пробелы схлопываются.
    assert reply.json()["positions"][0]["note"] == "дороже, но успеет к сроку"


def test_fail_otvet_pokazyvayet_novyy_sostav_srazu(client: Any, db: Any) -> None:
    """Ответ на действие показывает уже новый состав вариантов.

    Страница рисуется тем, что вернул сервер. Ответ с прежним составом
    означает, что человек нажал кнопку, ничего не изменилось, и он жмёт ещё
    раз — а потом обновляет страницу руками и обнаруживает, что оба нажатия
    сработали.
    """
    from platform_api.db.models import Role
    from tests.conftest import sign_in

    org = sign_in(db, client, Role.ANALYST)
    work = _work(db, org)
    db.commit()
    position_value = work.positions[0]
    was = len(position_value.options)

    reply = client.post(
        f"/api/tender/works/{work.id}/positions/{position_value.id}/ask",
        json={"name": "Пикобур PDC Ø215, 3 лопасти"},
    )
    assert reply.status_code == 200

    line = reply.json()["positions"][0]
    kinds = [option["source"] for option in line["options"]]
    assert "asked" in kinds, "заявки нет в ответе — состав отдан прежний"
    assert len(line["options"]) != was or kinds != ["found"] * was

    # И то же самое при подтверждении: отвергнутые уходят сразу, а не после F5.
    second = work.positions[1]
    reply = client.post(f"/api/tender/works/{work.id}/options/{second.options[0].id}/choose")
    choice = reply.json()["positions"][1]["options"]
    assert len(choice) == 1 and choice[0]["chosen"]

    # И при удалении: убранный вариант не должен вернуться в том же ответе.
    reply = client.delete(f"/api/tender/works/{work.id}/options/{choice[0]['id']}")
    assert reply.status_code == 200
    assert reply.json()["positions"][1]["options"] == []


def test_fail_zakaz_poiska_ubirayet_i_podtverzhdyonnuyu_nahodku(db: Any, organization: Any) -> None:
    """Заказ поиска убирает и подтверждённую находку модели.

    «Найдите сами» и «этот поставщик подтверждён» — противоречащие друг другу
    указания, и оставленные рядом они означают, что снабжение проверяет одного,
    пока разбор ждёт другого. Подтверждение — не сохранённая ценность: находку
    модели можно найти заново, а вот добавленное руками остаётся.
    """
    from platform_api.db.models import OptionSource
    from platform_api.modules.tender import works

    work = _work(db, organization)
    position_value = work.positions[0]
    works.choose(db, work, position_value.options[0].id)
    assert any(option.chosen for option in position_value.options)

    works.ask(db, work, position_value.id, "Пикобур PDC Ø215")

    kinds = {option.source for option in position_value.options}
    assert kinds == {OptionSource.ASKED}
    assert not any(option.chosen for option in position_value.options)


def test_fail_zadanie_ne_nazyvayet_zakazchika_i_mesta() -> None:
    """Задание не называет ни заказчика, ни куда везти.

    Снабжение пересылает его поставщику. Узнав заказчика и рудник, поставщик
    находит закупку в реестре, видит цену заключения — и либо поднимает свою
    до неё, либо идёт на торги сам. Это единственное, чем мы отличаемся от
    него в этой сделке: купить у завода он умеет не хуже.
    """
    from platform_api.modules.tender import spec

    raw_text = (
        "УТВЕРЖДАЮ\n"
        "Директор департамента Қайыргелды Б.К.\n"
        "Насос центробежный, напор не менее 150 м\n"
        "Место поставки: рудник «Западный Мынкудук», Туркестанская область\n"
        "Заказчик: ТОО «Семизбай-U»\n"
        "Поставка осуществляется в адрес ТОО «Семизбай-U» по заявке\n"
        "Контакты: zakup@kazatomprom.kz, +7 (7172) 45-90-58\n"
        "Закупка № 12345678-1\n"
    )
    cleaned = spec.limited(raw_text)

    assert "напор не менее 150 м" in cleaned
    for secret in (
        "Мынкудук",
        "Туркестанская",
        "Семизбай",
        "Қайыргелды",
        "kazatomprom.kz",
        "45-90-58",
        "12345678-1",
        "УТВЕРЖДАЮ",
    ):
        assert secret not in cleaned, f"в задании осталось: {secret}"


def test_fail_oblast_primeneniya_ne_prinimayut_za_adres() -> None:
    """«Область применения» — заголовок задания, а не адрес.

    Правило по одному слову «область» унесло бы начало половины технических
    заданий. Областью считается только «Туркестанская область» — с названием
    перед ней.
    """
    from platform_api.modules.tender import spec

    cleaned = spec.limited(
        "1. Область применения\nДиапазон измерений в области низких давлений\n"
        "Кызылординская область, город Байконур"
    )

    assert "Область применения" in cleaned and "области низких давлений" in cleaned
    assert "Кызылординская" not in cleaned


def test_fail_organizaciya_v_trebovanii_zamenyayetsya_a_ne_stirayet_stroku() -> None:
    """Организация посреди требования заменяется, а не уносит требование.

    «Согласовать с ТОО «Икс» до отгрузки» — это требование к поставке.
    Вычеркнутое целиком, оно исчезает вместе с обязанностью согласовывать.
    """
    from platform_api.modules.tender import spec

    cleaned = spec.limited("Отгрузка согласовывается с ТОО «Семизбай-U» за 5 дней")

    assert "Отгрузка согласовывается" in cleaned and "за 5 дней" in cleaned
    assert "Семизбай" not in cleaned


def test_fail_pravki_razbora_chistyatsya_pri_otpravke(db: Any, organization: Any) -> None:
    """Дописанное разбором вычищается в момент отправки.

    Собранное платформой задание обезличено с самого начала, но разбор правит
    его руками и вполне может вписать «доставка в Кызылорду» — по делу для
    себя и во вред в письме поставщику.
    """
    from platform_api.modules.tender import works

    work = _work(db, organization)
    for position_value in work.positions:
        works.choose(db, work, position_value.options[0].id)
    works.set_spec(
        db,
        work,
        work.positions[0].id,
        "ТЕХНИЧЕСКОЕ ЗАДАНИЕ\n\nПикобур PDC\nМесто поставки: рудник Каратау\nЗамок З-88",
    )
    assert "Каратау" in work.positions[0].spec

    works.hand_over(db, work, "проверьте цены")

    assert "Каратау" not in work.positions[0].spec
    assert "Замок З-88" in work.positions[0].spec


def test_fail_snabzhenie_ne_vidit_zakazchika(db: Any, organization: Any) -> None:
    """Заказчик не уходит снабжению и в шапке лота.

    Обезличенное задание рядом с именем заказчика на той же странице — это
    незакрытая дверь с табличкой «закрыто»: снабженец видит и то и другое.
    """
    from platform_api.db.models import Role
    from platform_api.modules.tender.works_router import _work_out

    work = _work(db, organization)

    assert _work_out(work, Role.ANALYST).customer == "АО «Волковгеология»"
    for_supply = _work_out(work, Role.BUYER)
    assert for_supply.customer == ""
    # Имя исходного файла тоже говорящее: «ТЗ ПНА. Западный Мынкудук.pdf».
    assert all(position.spec_source == "" for position in for_supply.positions)


def test_fail_zadanie_po_lotu_odnim_faylom(client: Any, db: Any) -> None:
    """По лоту — один файл со всеми позициями, а не шесть.

    Поставщику пишут одно письмо. Шесть вложений он прочтёт как шесть разных
    запросов и ответит на первое.
    """
    from io import BytesIO

    from docx import Document
    from platform_api.db.models import Role
    from tests.conftest import sign_in

    org = sign_in(db, client, Role.ANALYST)
    work = _work(db, org)
    db.commit()

    reply = client.get(f"/api/tender/works/{work.id}/spec.docx")
    assert reply.status_code == 200
    assert "filename*=UTF-8''" in reply.headers["content-disposition"]

    body = "\n".join(paragraph.text for paragraph in Document(BytesIO(reply.content)).paragraphs)
    for position_value in work.positions:
        assert position_value.title in body


def test_fail_imya_zakazchika_vychishchayetsya_pricelno(db: Any, organization: Any) -> None:
    """Заказчика этого лота платформа знает и убирает точно.

    Общие правила закрывают то, чего мы заранее не знаем. Имя заказчика мы
    знаем: оно записано в самой работе, и в документе встречается во всех
    написаниях сразу — «ТОО «Волковгеология»», «Волковгеология», «АО
    Волковгеология». Гадать по написанию, имея точное имя, незачем.
    """
    from platform_api.modules.tender import works

    work = _work(db, organization)
    assert work.customer == "АО «Волковгеология»"
    for position_value in work.positions:
        works.choose(db, work, position_value.options[0].id)
    works.set_spec(
        db,
        work,
        work.positions[0].id,
        "ТЗ\n\nПикобур PDC Ø190,5\nПоставка по заявке Волковгеология, замок З-88",
    )

    works.hand_over(db, work, "")

    brief = work.positions[0].spec
    assert "Волковгеология" not in brief
    # Товар при этом на месте: вычищается заказчик, а не наименование.
    assert "Пикобур PDC Ø190,5" in brief and "замок З-88" in brief


def test_fail_nazvaniye_tovara_perezhivayet_vychistku(db: Any, organization: Any) -> None:
    """Наименование позиции переживает вычистку.

    Соблазн вычистить заодно и название лота стоит дорого: это и есть товар, и
    задание без него — лист бумаги с требованиями неизвестно к чему.
    """
    from platform_api.modules.tender import works

    work = _work(db, organization)
    for position_value in work.positions:
        works.choose(db, work, position_value.options[0].id)
    works.hand_over(db, work, "")

    for position_value in work.positions:
        assert position_value.title.split()[0] in position_value.spec


def _row_stub(dir_path: str, sheet_name: str) -> Any:
    """Строка отбора настолько, насколько её видит раскладка по лотам."""
    from types import SimpleNamespace

    return SimpleNamespace(row=SimpleNamespace(folder_path=dir_path, title=sheet_name))


def test_fail_vzyatoye_v_rabotu_uhodit_vniz_spiska() -> None:
    """Взятое в работу опускается в конец отбора.

    Решение по нему принято, а место наверху занято тем, что уже не выбирают.
    Наверху должно стоять то, по чему ещё предстоит решить.
    """
    from platform_api.modules.tender.router import _grouped
    from platform_api.modules.tender.works import Taken

    line = _row_stub
    rows = [line("/архив/А", "Насос"), line("/архив/Б", "Кран"), line("/архив/В", "Задвижка")]
    taken = {("/архив/А", "Насос"): Taken(id="1", code="TN-00001", stage="analysis")}

    assert [item.row.title for item in _grouped(rows, {}, taken)] == [
        "Кран",
        "Задвижка",
        "Насос",
    ]


def test_fail_lot_uhodit_vniz_celikom() -> None:
    """Лот опускается вниз целиком, а не той позицией, которую взяли.

    Половина лота наверху и половина внизу — это два лота на вид, и вторую
    половину возьмут в работу ещё раз.
    """
    from platform_api.modules.tender.router import _grouped
    from platform_api.modules.tender.works import Taken

    rows = [
        _row_stub("/архив/А", "Пикобур Ø190"),
        _row_stub("/архив/Б", "Кран"),
        _row_stub("/архив/А", "Пикобур Ø215"),
    ]
    lot = {("/архив/А", "Пикобур Ø190"): "лот-1", ("/архив/А", "Пикобур Ø215"): "лот-1"}
    taken = {("/архив/А", "Пикобур Ø190"): Taken(id="1", code="TN-00001", stage="supply")}

    assert [item.row.title for item in _grouped(rows, lot, taken)] == [
        "Кран",
        "Пикобур Ø190",
        "Пикобур Ø215",
    ]


def test_fail_lot_ubirayet_tolko_administrator(client: Any, db: Any) -> None:
    """Убрать лот из работы может только администратор.

    Вместе с лотом уходят подтверждённые поставщики, найденные снабжением цены
    и правки заданий — работа двух отделов за неделю. Кнопка, доступная всем,
    однажды будет нажата вместо соседней.
    """
    import uuid as _uuid

    from platform_api.auth import passwords
    from platform_api.auth.service import open_session
    from platform_api.config import Settings
    from platform_api.db.models import (
        Membership,
        Role,
        TenderWork,
        TenderWorkOption,
        User,
    )
    from sqlalchemy import func, select
    from tests.conftest import sign_in

    org = sign_in(db, client, Role.ANALYST)
    work = _work(db, org)
    db.commit()
    # Имя запоминаем заранее: после удаления обращение к объекту поднимет
    # ObjectDeletedError, и тест упадёт на проверке вместо проверки.
    work_id = work.id

    assert client.delete(f"/api/tender/works/{work_id}").status_code == 403

    admin = User(
        email=f"{_uuid.uuid4().hex[:8]}@fintend.kz",
        password_hash=passwords.hash_password("закупки-2026-каратау"),
    )
    db.add(admin)
    db.flush()
    db.add(Membership(user_id=admin.id, organization_id=org.id, role=Role.ADMIN))
    db.flush()
    _, token = open_session(db, admin, org, ttl_hours=12)
    db.commit()
    client.cookies.set(Settings().auth.session_cookie, token)

    assert client.delete(f"/api/tender/works/{work_id}").status_code == 204

    # Насовсем, вместе с позициями и вариантами: смысл в том, чтобы взять лот
    # заново, а сохранённая работа заняла бы его код.
    db.expire_all()
    assert (
        db.execute(select(TenderWork).where(TenderWork.id == work_id)).scalar_one_or_none() is None
    )
    assert db.execute(select(func.count()).select_from(TenderWorkOption)).scalar() == 0


def test_fail_perenesyonnaya_stroka_sklеivayetsya() -> None:
    """Требование, разорванное переносом, склеивается до отбора.

    Разбор pdf отдаёт текст так, как он лёг на страницу: «Наработка на отказ
    не менее» на одной строке, «8000 часов» на следующей. Отбор по строкам на
    таком тексте выбрасывает половину требования и оставляет вторую — читать
    это невозможно, а искать по этому товар тем более.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ.pdf",
            text="Наработка на отказ не менее\n8000 часов с даты ввода\nв эксплуатацию",
        )
    )
    body = spec.gather(case).body

    assert "не менее 8000 часов с даты ввода в эксплуатацию" in body


def test_fail_povtoryonnyy_abzac_pechatayetsya_odin_raz() -> None:
    """Одно и то же требование печатается один раз.

    В техническом задании и приложении к нему оно стоит слово в слово, а
    таблица позиций повторяется на каждой странице. Задание, где каждый абзац
    написан дважды, перестают читать на третьем.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ.pdf",
            text="Степень защиты: IP 68.\nСтепень защиты: IP 68.\nМатериал: сталь 1.4401.",
        )
    )
    body = spec.gather(case).body

    assert body.count("IP 68") == 1 and "сталь 1.4401" in body


def test_fail_kazahskiy_dubl_ne_udvaivayet_zadaniye() -> None:
    """Казахская половина документа в задание не идёт.

    Техническое задание пишут на двух языках, и вторая половина дословно
    повторяет первую. Снабжение читает по-русски, а двойная длина означает,
    что не дочитают ни ту ни другую.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ.pdf",
            text=(
                "Наружный диаметр, мм - 190,5\n"
                "Сыртқы диаметрі, мм - 190,5 бұрғылау үшін\n"
                "Степень защиты: IP 68"
            ),
        )
    )
    body = spec.gather(case).body

    assert "Наружный диаметр" in body and "IP 68" in body
    assert "Сыртқы" not in body


def test_fail_shapka_bez_strok_ne_pechatayetsya() -> None:
    """Таблица, из которой вырезали все данные, не оставляет шапки.

    Из списка поставщиков заказчика уходят и цены, и названия — остаётся
    «№ · Наименование компании». Такой остов говорит только о том, что таблица
    была, а читается как обрывок.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "МЗ",
            "МЗ.docx",
            text="Степень защиты: IP 68",
            tables=[
                (
                    ("№", "Наименование компании", "Статус"),
                    ("1", "ТОО «Ромашка»", "Дистрибьютор"),
                    ("2", "ТОО «Роза»", "Поставщик"),
                )
            ],
        )
    )
    body = spec.gather(case).body

    assert "IP 68" in body
    assert "Наименование компании" not in body and "Ромашка" not in body


def test_fail_chuzhaya_poziciya_v_zadaniye_ne_popadayet() -> None:
    """В задании на кабель не должно быть индукционного котла.

    В папке одна закупка на шестнадцать позиций и одно техническое задание на
    всех. Приложенное целиком к каждой, оно отправляет снабжению задание, в
    котором описан соседний товар, — и снабженец ищет не то.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ.pdf",
            text=(
                "Кабель КГ 3х50 должен соответствовать ГОСТ 24334-80, "
                "число жил не менее трёх, сечение 50 мм2, изоляция резиновая, "
                "температура эксплуатации не ниже минус 40 градусов.\n"
                "Индукционный котёл должен иметь мощность не менее 30 кВт, "
                "напряжение 380 В, класс защиты IP 54,материал корпуса сталь, "
                "давление в контуре не более 0,6 МПа, масса не более 40 кг.\n"
                "Сертификат качества обязателен для всех позиций закупки, "
                "паспорт изделия на государственном и русском языках прилагается."
            ),
        )
    )
    case.requested = (
        type(case.requested[0])(
            name="Кабель КГ 3х50",
            specification=None,
            quantity=Decimal(1),
            unit="м",
            ens_code="",
        ),
        type(case.requested[0])(
            name="Индукционный котёл ВИН 30",
            specification=None,
            quantity=Decimal(1),
            unit="шт",
            ens_code="",
        ),
    )
    текст = spec.render(spec.gather(case), "Кабель КГ 3х50")

    assert "ГОСТ 24334-80" in текст
    assert "Индукционный" not in текст
    # Общее требование остаётся: чужого имени в нём нет.
    assert "Сертификат качества" in текст


def test_fail_usloviya_oplaty_ne_uhodyat() -> None:
    """Как мы рассчитываемся с заказчиком — не дело поставщика.

    «Окончательный платёж 100%» в письме поставщику читается как наше
    предложение ему, и торг начинается с этой строки.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "МЗ",
            "МЗ.docx",
            text=(
                "Окончательный платеж - 100%, промежуточный - 0%\n"
                "Предоплата - 0%\n"
                "Объём не менее 300 л, с морозильным отделом"
            ),
        )
    )
    body = spec.gather(case).body

    assert "не менее 300 л" in body
    assert "платеж" not in body.lower() and "Предоплата" not in body


def test_fail_poziciya_beryot_svoy_dokument_i_tehzadaniye() -> None:
    """Позиция берёт свой документ и техническое задание закупки.

    В заключении позиция только перечислена строкой таблицы, а описана она в
    техническом задании той же папки: диаметр, резцы, ГОСТ. Взяв одно
    заключение, мы отправляем снабжению строку вместо требований; взяв одно
    задание — рискуем описанием соседней нужды из чужой записки.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "ТЗ насосы.pdf",
            text="Материал корпуса: нержавеющая сталь 1.4401. Степень защиты: IP 68.",
        ),
        _document(
            "МЗ",
            "Проект МЗ.docx",
            text="Насос ЦНС-60 должен поставляться в заводской упаковке",
        ),
    )
    case.requested[0].source_document = "Проект МЗ.docx"

    текст = spec.render(spec.gather(case), "Насос ЦНС-60")

    assert "заводской упаковке" in текст, "потерян свой документ позиции"
    assert "1.4401" in текст and "IP 68" in текст, "потеряно техническое задание"


def test_fail_stroka_chuzhogo_plana_zakupok_ne_popadayet() -> None:
    """Строка чужого товара из плана закупок в задание не идёт.

    План перечисляет весь месяц одной таблицей, и его товары в наши позиции не
    входят вовсе — по именам позиций их не опознать. Поэтому товарная строка
    обязана назвать свою позицию: не назвала — она про что-то другое.

    Строка характеристики при этом остаётся: «Напор · 150 м» своего имени не
    содержит и содержать не должно.
    """
    from platform_api.modules.tender import spec

    case = _case_with_documents(
        _document(
            "ТЗ",
            "План.xlsx",
            text="Напор не менее 150 м",
            tables=[
                (
                    ("№", "Наименование", "Код", "Ед.", "Кол-во", "Доставка"),
                    ("1", "Насос ЦНС-60", "281314", "шт", "3", "DDP"),
                    ("2", "Кран шаровой ППР Ду15", "275021", "шт", "12", "DDP"),
                ),
                (("Параметр", "Значение"), ("Напор", "150 м")),
            ],
        )
    )
    case.requested = (
        case.requested[0],
        type(case.requested[0])(
            name="Кран шаровой ППР Ду15",
            specification=None,
            quantity=Decimal(12),
            unit="шт",
            ens_code="",
            source_document="План.xlsx",
        ),
    )
    текст = spec.render(spec.gather(case), "Насос ЦНС-60")

    assert "Насос ЦНС-60 · 281314" in текст
    assert "Кран шаровой" not in текст
    assert "Напор · 150 м" in текст
