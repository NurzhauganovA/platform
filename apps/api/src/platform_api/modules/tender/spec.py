"""Техническое задание позиции — то единственное, что видит снабжение.

Снабжению нужен предмет закупки и требования к нему. Исходный документ
заказчика ему открывать нельзя: в ТЗ стоят цены ценового заключения, реквизиты
сторон и печати, а от цены заказчика считается наша — увидев её, снабженец
знает потолок, по которому с ним же и будут торговаться поставщики.

Поэтому задание не вычищается, а **собирается заново** — из того, что ядро уже
разобрало, и только из перечисленных здесь полей. Разница принципиальная:
вычистка отвечает «мы убрали всё, что вспомнили», сборка — «сюда попало только
то, что названо». Первая ошибается молча и ровно один раз.

Платного шага здесь нет. Всё, из чего собирается задание, ядро извлекло при
разборе документов, и лежит это в его базе: открытая страница не тратит денег.

Черновик записывается в саму работу при взятии её в работу, а не считается на
каждый показ. Отбор пересобирается при каждом прогоне, названия у позиций
меняются — а задание, отданное снабжению, меняться от этого не должно.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from decimal import Decimal
from typing import Any

MAX_LENGTH = 40_000
"""Предел длины задания.

Половина технических заданий укладывается в две тысячи знаков, девять из
десяти — в тринадцать тысяч. Сорок тысяч закрывают почти всё; у остатка хвост
отрезается с прямой пометкой об этом. Молча обрезанное задание хуже длинного:
снабжение не знает, что читает половину.
"""


@dataclass(frozen=True, slots=True)
class Position:
    """Позиция глазами заказчика. Денежных полей здесь нет намеренно."""

    name: str
    specification: str
    quantity: Decimal | None
    unit: str
    ens_code: str


@dataclass(frozen=True, slots=True)
class Source:
    """Документ заказчика, из которого берётся задание.

    Хранится текстом, а не ссылкой на файл: исходный документ снабжению
    закрыт, а текст его технической части — ровно то, что ему нужно.
    """

    name: str
    kind: str
    body: str


@dataclass(frozen=True, slots=True)
class CaseSpec:
    """Требования закупки — белый список полей, годных для снабжения.

    Собирается попутно со строками отбора: закупка в этот момент уже построена,
    и отдельного обращения к базе ядра это не стоит.
    """

    subject: str
    requirements: tuple[str, ...]
    delivery: str
    warranty: str
    positions: tuple[Position, ...]
    source: str
    """Из какого документа собрано. Спорное требование нужно уметь проверить —
    разбору исходный документ по-прежнему открыт."""

    body: str = ""
    """Техническая часть документа заказчика — его собственный текст.

    Пересказ модели годился бы, если бы речь шла о понимании. Но снабжение по
    этому тексту закупает: марка стали, допуск, класс защиты и номер ГОСТа
    должны стоять теми же словами, что и у заказчика. Выжимка в три строки
    теряет как раз таблицу с размерами.
    """


def gather(case: Any) -> CaseSpec:
    """Требования закупки из разобранных документов.

    Требования берутся из ТЗ и МЗ — документов заказчика. Из коммерческих
    предложений их брать нельзя: там поставщик пишет, что удобно ему.
    """
    заказчик = [
        document.insight
        for document in case.documents
        if document.insight is not None and str(document.insight.kind) in _CUSTOMER_KINDS
    ]
    источник = next(
        (
            document.source.name
            for document in case.documents
            if document.insight is not None and str(document.insight.kind) == "ТЗ"
        ),
        "",
    )
    документ = _customer_document(case)
    return CaseSpec(
        subject=_line(case.subject),
        requirements=_requirements(заказчик, case),
        delivery=_first(insight.delivery_terms for insight in заказчик),
        warranty=_first(insight.warranty for insight in заказчик),
        positions=tuple(
            Position(
                name=_line(item.name),
                specification=_line(item.specification),
                quantity=item.quantity,
                unit=_line(item.unit),
                ens_code=_line(item.ens_code),
            )
            for item in case.requested
        ),
        source=документ.name or источник,
        body=документ.body,
    )


_CUSTOMER_KINDS = frozenset({"ТЗ", "МЗ"})
"""Чьи требования считаем требованиями закупки. Документы поставщиков сюда не
входят: в КП написано то, что удобно приславшему."""


def draft(folder_path: str, title: str) -> tuple[str, str]:
    """Черновик задания по позиции: сам текст и имя документа-источника.

    Пусто, если закупку ещё не разбирали. Пустое задание честнее выдуманного:
    снабжение по нему поймёт, что задание нужно написать руками, а не что
    требований к товару нет.
    """
    from platform_api.modules.tender import worklist

    known = worklist.case_spec(folder_path)
    if known is None:
        return "", ""
    return render(known, title), known.source


def render(spec: CaseSpec, title: str) -> str:
    """Задание по позиции — текстом, каким его и увидят.

    Текстом, а не набором полей: задание правит человек из разбора перед
    передачей, и «дописать строчку» должно оставаться дописыванием строчки.
    Разложенное по восьми полям задание правят так же часто, как заполняют
    восемь полей, то есть никогда.
    """
    выбрана = _match(spec.positions, title)
    позиции = (выбрана,) if выбрана is not None else spec.positions
    if not позиции:
        # Позиций заказчика ядро не нашло — папка без ценового заключения.
        # Название строки здесь единственное, что известно о товаре, и без
        # него задание выходит без предмета вовсе.
        позиции = (
            Position(name=_line(title), specification="", quantity=None, unit="", ens_code=""),
        )

    части: list[str] = ["ТЕХНИЧЕСКОЕ ЗАДАНИЕ", ""]
    if spec.subject and (выбрана is None or _same(spec.subject) != _same(выбрана.name)):
        части += [f"Предмет закупки: {spec.subject}", ""]

    for index, position in enumerate(позиции, start=1):
        номер = f"{index}. " if len(позиции) > 1 else ""
        части.append(f"{номер}{position.name or _line(title)}")
        if position.quantity is not None:
            части.append(f"   Количество: {_amount(position.quantity)} {position.unit}".rstrip())
        if position.ens_code:
            части.append(f"   Код ЕНС ТРУ: {position.ens_code}")
        if position.specification:
            части.append(f"   Характеристики: {position.specification}")
        части.append("")

    if spec.body:
        # Текст самого документа заказчика. Пересказ здесь не годится: марка
        # стали, допуск и номер ГОСТа должны стоять теми же словами, иначе
        # снабжение купит похожее.
        части += [f"ИЗ ДОКУМЕНТА ЗАКАЗЧИКА · {spec.source}".rstrip(" ·"), "", spec.body]
    else:
        # Документ не прочитался — собираем из того, что ядро из него поняло.
        if spec.requirements:
            части.append("ТРЕБОВАНИЯ ЗАКАЗЧИКА")
            части += [f"— {item}" for item in spec.requirements]
            части.append("")
        if spec.delivery:
            части += ["МЕСТО И УСЛОВИЯ ПОСТАВКИ", spec.delivery, ""]
        if spec.warranty:
            части += ["ГАРАНТИЯ", spec.warranty, ""]

    готово = "\n".join(части).strip()
    if len(готово) <= MAX_LENGTH:
        return готово
    # Обрезанное молча задание хуже длинного: снабжение не узнает, что читает
    # половину, и закупит по половине.
    return готово[:MAX_LENGTH] + "\n\n[…] Документ длиннее — запросите остаток у отдела разбора"


def document(title: str, code: str, body: str) -> bytes:
    """Задание файлом .docx — тем, что открывается у поставщика и печатается.

    Снабжение пересылает задание наружу: поставщику нужно понять, что искать.
    Пересылать ссылку на нашу платформу нельзя — там за той же позицией стоит
    себестоимость, а «просто не заходите в соседнюю вкладку» отделом снабжения
    не обеспечивается.
    """
    from io import BytesIO

    from docx import Document
    from docx.shared import Pt

    файл = Document()
    обычный = файл.styles["Normal"]
    обычный.font.name = "Calibri"
    обычный.font.size = Pt(11)

    заголовок = файл.add_paragraph()
    прогон = заголовок.add_run(f"{code} · {title}" if code else title)
    прогон.bold = True
    прогон.font.size = Pt(13)

    for строка in body.splitlines():
        абзац = файл.add_paragraph(строка)
        # Заголовки разделов набраны прописными — по ним и отбиваются, чтобы
        # не заводить вторую разметку рядом с той, которую правит человек.
        if строка and строка == строка.upper() and any(знак.isalpha() for знак in строка):
            абзац.runs[0].bold = True

    поток = BytesIO()
    файл.save(поток)
    return поток.getvalue()


# ---------------------------------------------------------------------------


_MONEY = re.compile(
    r"тенге|₸|\bkzt\b|\bндс\b|стоимост|\bцен[аыуеой]|\bцен\b|сумм|оплат|предоплат|аванс",
    re.IGNORECASE,
)
"""Признаки денежного условия в требовании.

Требования заказчика идут в задание как есть, кроме коммерческих: «цена не
выше стольких-то» — это наш потолок, а не свойство товара. Техническое
требование такими словами почти не пишется, а цена, ушедшая снабжению, торгует
против нас.
"""


_CUSTOMER_ORDER = ("ТЗ", "МЗ")
"""Чей текст берём и в каком порядке.

Техническое задание — то, чему поставка должна соответствовать. Маркетинговое
заключение вторым: в нём то же требование, но короче, а бывает оно чаще —
закупка без ТЗ обычное дело.
"""


def _customer_document(case: Any) -> Source:
    """Техническая часть документа заказчика — его собственными словами.

    Текст уже разобран ядром и лежит в его базе: платить за чтение файла
    второй раз незачем, а открытая страница не тратит денег.
    """
    for kind in _CUSTOMER_ORDER:
        for document in case.documents:
            insight = document.insight
            if insight is None or str(insight.kind) != kind:
                continue
            текст = _readable(document.extraction)
            if текст:
                return Source(name=document.source.name, kind=kind, body=текст)
    return Source(name="", kind="", body="")


def _readable(extraction: Any) -> str:
    """Текст документа и его таблицы, без денежных строк и реквизитов.

    Таблицы идут наравне с текстом: в техническом задании размеры, допуски и
    количества стоят именно в них, и документ без таблиц — это документ без
    того, ради чего его читают.

    Отбор построчный, а не по документу целиком. Убрать страницу из-за одной
    цены значит убрать вместе с ней требования, ради которых страницу и
    открывали.
    """
    строки: list[str] = []
    for page in extraction.pages:
        строки += (page.text or "").splitlines()
        строки += [
            " · ".join(cell.strip() for cell in row if cell and cell.strip())
            for table in page.tables
            if table.is_meaningful
            for row in table.rows
            if not _priced_row(row)
        ]

    очищены: list[str] = []
    пусто = False
    for строка in строки:
        текст = " ".join(строка.split())
        if not текст:
            # Пустые строки сжимаются в одну: разбор pdf оставляет их пачками,
            # и задание превращается в лестницу.
            пусто = bool(очищены)
            continue
        if _commercial(текст) or _REQUISITES.search(текст) or _SIGNATURE.match(текст):
            continue
        if пусто:
            очищены.append("")
            пусто = False
        очищены.append(текст)
    return "\n".join(очищены).strip()


_CURRENCY = re.compile(r"тенге|₸|\bkzt\b|\bтг\b|\bндс\b", re.IGNORECASE)
"""Валюта названа прямо — строка про деньги, чем бы она ни была."""

_PRICE_WORD = re.compile(
    r"стоимост|\bсумм[аыуе]?\b|\bцен[аыуеой]\b|прайс|оплат|аванс", re.IGNORECASE
)
_BIG_NUMBER = re.compile(r"\d[\d\s\u00a0\u202f]{3,}")
"""Четыре цифры подряд и больше. Столько бывает у суммы и почти не бывает у
допуска: «±1 г/см³» и «не менее 8000 часов» проходят, «350 000» — нет."""

_REQUISITES = re.compile(
    r"\bбин\b|\bиин\b|\bиик\b|\bбик\b|\bкбе\b|\bкнп\b|реквизит|"
    r"расч[ёе]тный сч[ёе]т|\bм\.?\s?п\.|\bподпис[ьи]\b|\bпечат[ьи]\b|факсимиле",
    re.IGNORECASE,
)
"""Реквизиты и место печати. Снабжению они не нужны, а в задании, которое оно
перешлёт поставщику, выглядят как наши реквизиты."""


_AMOUNT = re.compile(r"^\d{1,3}(?:[\s\u00a0\u202f]\d{3})+[.,]\d{2}$|^\d{4,}[.,]\d{2}$")
"""Ячейка, записанная как денежная сумма: «1 000,00», «12500,00».

Отличать деньги от количества приходится по форме записи, а не по соседнему
слову: в ценовом заключении колонки цен подписаны один раз в шапке, а в
строках стоят голые числа. Разряды, разделённые пробелом, и два знака после
запятой — то, как в казахстанских документах пишут сумму и почти никогда не
пишут количество: «5,00 шт.» бывает, «1 000,00 шт.» — нет.
"""

_DECIMAL = re.compile(r"^\d+[.,]\d{2}$")

_SIGNATURE = re.compile(r"^[_\s]{3,}$|^«?_+»?\s*_+\s*20\d\d")
"""Место подписи и дата прописью — «_______ /___/», ««___» ______ 2026 г.».
Это и есть печать в бумажном смысле; в задании она означает, что снабжение
переслало поставщику бланк заказчика."""


def _priced_row(row: Any) -> bool:
    """Строка таблицы с ценами.

    Ценовое заключение — документ про цены, и его главная таблица это цены
    поставщиков по позиции. Одна такая строка, ушедшая снабжению, сообщает
    поставщику, во сколько заказчик оценил закупку и почём предлагали соседи.

    Строка позиции при этом остаётся: в ней количество «5,00», а не сумма
    «1 000,00», и по форме записи одно от другого отличается.
    """
    ячейки = [" ".join(str(cell).split()) for cell in row if cell]
    if any(_AMOUNT.match(cell) for cell in ячейки):
        return True
    return sum(1 for cell in ячейки if _DECIMAL.match(cell)) >= 2


_AMOUNT_INSIDE = re.compile(r"\d{1,3}(?:[\s\u00a0\u202f]\d{3})+[.,]\d{2}(?!\d)")
"""Сумма посреди строки: «785 000,00».

Таблицу цен разбор файла нередко отдаёт не таблицей, а строкой текста с
разделителями — и правило, применённое только к ячейкам, её пропускает. Форма
записи при этом та же, и ловится она там же.
"""


def _commercial(line: str) -> bool:
    """Строка про деньги.

    Одного слова «цена» мало: «цена деления 1 г/см³» — это характеристика
    ареометра, и выбросить её значит купить не тот прибор. Деньгами строку
    делает названная валюта, денежное слово рядом с крупным числом или сама
    форма записи суммы — разряды через пробел и два знака после запятой.

    Технические величины так не пишут: «1000 м³», «не менее 8000 часов», но не
    «1 000,00 м³». Хвост «,00» — привычка бухгалтерская, не инженерная.
    """
    if _CURRENCY.search(line) or _AMOUNT_INSIDE.search(line):
        return True
    if len(_DECIMAL_INSIDE.findall(line)) >= 2:
        # Две дроби с двумя знаками в одной строке — это колонки цен, даже
        # когда шапка с их названиями осталась страницей выше.
        return True
    return bool(_PRICE_WORD.search(line) and _BIG_NUMBER.search(line))


_DECIMAL_INSIDE = re.compile(r"(?<![\d.,])\d+[.,]\d{2}(?![\d.,])")


def _requirements(insights: list[Any], case: Any) -> tuple[str, ...]:
    """Требования заказчика без денежных условий, без повторов, по порядку.

    Один проход по объединению: требования повторяются в ТЗ и МЗ дословно, и
    без сведения задание начинается с трёх одинаковых абзацев.
    """
    источники = [item for insight in insights for item in insight.requirements] or list(
        case.requirements
    )
    видели: set[str] = set()
    отобраны: list[str] = []
    for item in источники:
        текст = _line(item)
        ключ = _same(текст)
        if not текст or ключ in видели or _MONEY.search(текст):
            continue
        видели.add(ключ)
        отобраны.append(текст)
    return tuple(отобраны)


def _match(positions: tuple[Position, ...], title: str) -> Position | None:
    """Позиция закупки, отвечающая строке отбора.

    Строка отбора берёт название прямо из позиции заказчика, поэтому сходится
    оно точно. Когда не сходится — закупка одной позиции, у строки название
    всей закупки: тогда позиция и так одна, и брать её можно без сравнения.
    """
    искомое = _same(title)
    точная = next((item for item in positions if _same(item.name) == искомое), None)
    if точная is not None:
        return точная
    return positions[0] if len(positions) == 1 else None


def _first(values: Any) -> str:
    return next((_line(value) for value in values if _line(value)), "")


def _line(value: Any) -> str:
    return " ".join(str(value).split()) if value else ""


def _same(text: str) -> str:
    return " ".join(text.casefold().split())


def _amount(value: Decimal) -> str:
    """Количество без хвоста нулей: «105», а не «105,000»."""
    целое = value.to_integral_value()
    return f"{целое:f}" if value == целое else f"{value.normalize():f}".replace(".", ",")


__all__ = ["MAX_LENGTH", "CaseSpec", "Position", "document", "draft", "gather", "render"]
